"""
WISDOM Lab - AI 리포트 → Word(.docx) 변환 모듈
제목 체계: Ⅰ.(14pt Bold) / 1.(13pt Bold) / (1)(12pt Bold) / 본문(11pt)
마크다운 표(| col | col |) 자동 변환
"""

import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 로마 숫자 대제목 패턴 (Ⅰ. Ⅱ. Ⅲ. …)
_RE_H1 = re.compile(r"^([ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ])\.\s+(.+)$")
# ── 아라비아 숫자 중제목 (1. 2. …)  ← 소제목 (1) 과 구분
_RE_H2 = re.compile(r"^(\d+)\.\s+(.+)$")
# ── 소제목 (1) (2) …
_RE_H3 = re.compile(r"^\((\d+)\)\s+(.+)$")
# ── 불릿 항목
_RE_BULLET = re.compile(r"^(\s*)([-*•]|\d+\.)\s+(.+)$")
# ── 마크다운 표 행
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[-| :]+\|$")


def _set_korean_font(run, size_pt: float, bold: bool = False,
                     color: RGBColor = None, italic: bool = False):
    """run에 한글/영문 폰트·크기·스타일 일괄 적용"""
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")


def _para_spacing(p, before_pt: float = 0, after_pt: float = 0, line_rule=None):
    """단락 간격 설정"""
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = Pt(18) if line_rule is None else line_rule


def _add_inline_runs(p, text: str, size_pt: float, bold_base: bool = False,
                     color: RGBColor = None):
    """**볼드** 인라인 마크업을 파싱해 run 분리 추가"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_korean_font(run, size_pt, bold=True, color=color)
        else:
            run = p.add_run(part)
            _set_korean_font(run, size_pt, bold=bold_base, color=color)


# ── 제목 추가 함수들 ─────────────────────────────────────────

def _add_h1(doc: Document, text: str):
    """Ⅰ. 대제목 — 14pt Bold, 진한 남색, 하단 구분선"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=16, after_pt=4)

    run = p.add_run(text)
    _set_korean_font(run, 14, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    # 하단 구분선
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_h2(doc: Document, text: str):
    """1. 중제목 — 13pt Bold, 파란색"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=10, after_pt=2)
    run = p.add_run(text)
    _set_korean_font(run, 13, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))
    return p


def _add_h3(doc: Document, text: str):
    """(1) 소제목 — 12pt Bold, 진회색"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=6, after_pt=2)
    run = p.add_run(text)
    _set_korean_font(run, 12, bold=True, color=RGBColor(0x40, 0x40, 0x40))
    return p


def _add_body(doc: Document, text: str):
    """본문 — 11pt, 줄간격 18pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_spacing(p, before_pt=0, after_pt=2)
    _add_inline_runs(p, text, 11)
    return p


def _add_bullet_item(doc: Document, text: str, indent_level: int = 0):
    """불릿 항목 — 11pt, 들여쓰기"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before_pt=0, after_pt=1)
    p.paragraph_format.left_indent = Cm(0.8 + indent_level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.5)

    # 불릿 기호
    bullet_run = p.add_run("▪ ")
    _set_korean_font(bullet_run, 11, bold=False, color=RGBColor(0x2E, 0x75, 0xB6))
    _add_inline_runs(p, text, 11)
    return p


# ── 마크다운 표 변환 ─────────────────────────────────────────

def _parse_table_rows(lines: list, start: int):
    """start 위치부터 표 행들을 파싱해 (header, body_rows, end_idx) 반환"""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if _RE_TABLE_SEP.match(line):
            i += 1
            continue
        if _RE_TABLE_ROW.match(line):
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
            i += 1
        else:
            break
    header = rows[0] if rows else []
    body = rows[1:] if len(rows) > 1 else []
    return header, body, i


def _add_table(doc: Document, header: list, body: list):
    """Word 표 생성"""
    if not header:
        return
    cols = len(header)

    # 테두리 설정 헬퍼
    def _set_cell_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "AAAAAA")
            tcBorders.append(el)
        tcPr.append(tcBorders)

    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.style = "Table Grid"

    # 표 전체 너비
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9026")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # 헤더 행
    hdr_row = table.rows[0]
    for ci, cell_text in enumerate(header):
        cell = hdr_row.cells[ci]
        cell.text = ""
        _set_cell_border(cell)
        # 헤더 배경
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6E4F0")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        _set_korean_font(run, 10.5, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    # 데이터 행
    for ri, row_data in enumerate(body):
        row = table.rows[ri + 1]
        for ci in range(cols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            cell = row.cells[ci]
            cell.text = ""
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text, 10.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    # 표 뒤 여백 단락
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── 메인 변환 함수 ───────────────────────────────────────────

def markdown_to_docx_bytes(
    md_text: str,
    title: str = "WISDOM Lab 분석 리포트",
    subject: str = "",
    doc_type_label: str = "연구보고서",
) -> bytes:
    """
    마크다운 리포트 텍스트 → Word(.docx) bytes 변환

    제목 체계:
      Ⅰ. → 14pt Bold (대제목)
      1.  → 13pt Bold (중제목)
      (1) → 12pt Bold (소제목)
      본문 → 11pt
    """
    doc = Document()

    # ── 페이지 설정 (A4, 여백 25mm) ──────────────────────────
    section = doc.sections[0]
    section.page_width  = 595 * 914  // 100   # ~11906 EMU → DXA 변환 필요 없음, Inches 사용
    section.page_height = 842 * 914  // 100
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── 기본 스타일 ────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(11)
    nrPr = normal.element.get_or_add_rPr()
    nrFonts = OxmlElement("w:rFonts")
    nrFonts.set(qn("w:eastAsia"), "맑은 고딕")
    nrPr.insert(0, nrFonts)

    # ── 표지 ──────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(cover, before_pt=40, after_pt=8)
    title_run = cover.add_run(title)
    _set_korean_font(title_run, 20, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    if subject:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(sp, before_pt=4, after_pt=4)
        sr = sp.add_run(f"주제: {subject}")
        _set_korean_font(sr, 12, color=RGBColor(0x44, 0x72, 0xC4))

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(tp, before_pt=2, after_pt=24)
    tr = tp.add_run(f"[{doc_type_label}]")
    _set_korean_font(tr, 11, italic=True, color=RGBColor(0x80, 0x80, 0x80))

    # 구분선
    sep = doc.add_paragraph()
    _para_spacing(sep, before_pt=0, after_pt=12)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single"); btm.set(qn("w:sz"), "6")
    btm.set(qn("w:space"), "1");    btm.set(qn("w:color"), "2E75B6")
    pBdr.append(btm); pPr.append(pBdr)

    # ── 본문 파싱 ─────────────────────────────────────────────
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        # 표 시작 감지
        if _RE_TABLE_ROW.match(line):
            header, body, next_i = _parse_table_rows(lines, i)
            _add_table(doc, header, body)
            i = next_i
            continue

        # 표 구분자 단독 등장 (이미 소비됐어야 하지만 안전 처리)
        if _RE_TABLE_SEP.match(line):
            i += 1
            continue

        # 대제목 Ⅰ.
        m = _RE_H1.match(line)
        if m:
            _add_h1(doc, line)
            i += 1
            continue

        # 중제목 1.
        m = _RE_H2.match(line)
        if m:
            _add_h2(doc, line)
            i += 1
            continue

        # 소제목 (1)
        m = _RE_H3.match(line)
        if m:
            _add_h3(doc, line)
            i += 1
            continue

        # 마크다운 ## / # 제목 (AI가 혼용하는 경우 대비)
        if line.startswith("### "):
            _add_h3(doc, line[4:].strip())
            i += 1
            continue
        if line.startswith("## ") and not line.startswith("### "):
            _add_h2(doc, line[3:].strip())
            i += 1
            continue
        if line.startswith("# ") and not line.startswith("## "):
            _add_h1(doc, line[2:].strip())
            i += 1
            continue

        # 불릿
        m = _RE_BULLET.match(raw)
        if m:
            indent = len(m.group(1)) // 2
            _add_bullet_item(doc, m.group(3).strip(), indent)
            i += 1
            continue

        # 구분선 ---
        if re.match(r"^[-_*]{3,}$", line):
            sep2 = doc.add_paragraph()
            _para_spacing(sep2, before_pt=4, after_pt=4)
            pPr2 = sep2._p.get_or_add_pPr()
            pBdr2 = OxmlElement("w:pBdr")
            b2 = OxmlElement("w:bottom")
            b2.set(qn("w:val"), "single"); b2.set(qn("w:sz"), "4")
            b2.set(qn("w:space"), "1");    b2.set(qn("w:color"), "CCCCCC")
            pBdr2.append(b2); pPr2.append(pBdr2)
            i += 1
            continue

        # 빈 줄 → 건너뜀 (문단 사이 공백 없음 규칙)
        if not line:
            i += 1
            continue

        # 일반 본문
        _add_body(doc, line)
        i += 1

    # ── 푸터 ──────────────────────────────────────────────────
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("WISDOM Lab  ·  Generated by AI")
    _set_korean_font(fr, 8, color=RGBColor(0xAA, 0xAA, 0xAA))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
